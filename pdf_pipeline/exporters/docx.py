"""DOCX export module"""

from docx import Document
from docx.shared import Inches
from typing import Dict, Any, List


def export_to_docx(document: Dict[str, Any], output_path: str):
    """
    Export the processed document to DOCX format
    
    Args:
        document: Processed document dictionary
        output_path: Path to save the DOCX file
    """
    doc = Document()
    
    # Add title
    title = document.get("metadata", {}).get("title", "Processed PDF Document")
    doc.add_heading(title, 0)
    
    # Process each element in the document structure
    for element in document["structure"]:
        elem_type = element["type"]
        
        if elem_type == "text":
            content = element.get("content", "")
            if content.strip():
                doc.add_paragraph(content)
        
        elif elem_type == "table":
            # Add table to document
            table_data = element
            add_table_to_docx(doc, table_data)
        
        elif elem_type == "image_caption":
            caption = element.get("content", "")
            if caption.strip():
                doc.add_paragraph(caption, style='Caption')
    
    # Save the document
    doc.save(output_path)


def add_table_to_docx(doc, table_data: Dict[str, Any]):
    """
    Add a table to the DOCX document
    
    Args:
        doc: python-docx Document object
        table_data: Table data dictionary
    """
    rows_data = table_data.get("rows", [])
    cols_data = table_data.get("columns", [])
    
    if not rows_data and not cols_data:
        return
    
    # Determine number of rows and columns
    num_cols = len(cols_data) if cols_data else (len(rows_data[0]) if rows_data else 0)
    num_rows = len(rows_data) + (1 if cols_data else 0)  # +1 for header if columns exist
    
    if num_cols == 0 or num_rows == 0:
        return
    
    # Create table in document
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    
    row_idx = 0
    
    # Add header row if columns are provided
    if cols_data:
        hdr_cells = table.rows[row_idx].cells
        for i, col_name in enumerate(cols_data):
            if i < num_cols:
                hdr_cells[i].text = str(col_name)
        row_idx += 1
    
    # Add data rows
    for data_row in rows_data:
        if row_idx >= num_rows:
            break
        row_cells = table.rows[row_idx].cells
        for i, cell_data in enumerate(data_row):
            if i < num_cols:
                row_cells[i].text = str(cell_data)
        row_idx += 1


def format_document_styles(doc):
    """
    Apply formatting styles to the document
    
    Args:
        doc: python-docx Document object
    """
    # Define styles if they don't exist
    styles = doc.styles
    
    # You can customize styles here as needed
    pass